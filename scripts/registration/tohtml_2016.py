# -*- coding: utf-8 -*-
import csv
import sys
import time
import os

from docx import Document
from docx.shared import Pt

document = Document()

from collections import OrderedDict

registered_email         = OrderedDict()
registered_fullname      = OrderedDict()
conf_year                = 2016
conf_place               = "Jeju, South Korea"
main_registrations_csv   = 'Main2016.csv'   #  All reciepts exported, e.g. Export-OCNS-Receipts-475-25-Jun-2015-05-35-34.csv
add_to_registrations_csv = 'AddToReg.csv'   #  Add to regs..
extras_csv               = 'Extras.csv'   #  Extras..
display                  = True
countries                = {}

tshirts = {'S'  : 0, 
           'M'  : 0, 
           'L'  : 0,
           'XL' : 0}

meeting = {'main'     : 0,
           'tutorial' : 0,
           'ws1'      : 0,
           'ws2'      : 0}

extras  = {'banquet' : 0}

members = {'faculty' : 0,
           'postdoc' : 0,
           'student' : 0}

non_members = {'faculty' : 0,
               'postdoc' : 0,
               'student' : 0}

to_write_full = {"First name" : 'first_name',
                 "Middle"     : 'middle_name',
                 "Surname"    : 'last_name',
                 "E-mail"     : 'email',
                 "Institution" : 'inst', 
                 "City"        : 'city',
                 "Country"     : 'country',
                 "Member type" : "type0", 
                 "Type of registration" : "paid",
                 "Banquet" : "banquet",
                 "Special dietary requirements" : "meal",
                 "T-Shirts" : "shirt"}

to_write_badges = {"Name" : 'full_name', 
                    "Member Type" : 'type0', 
                    "Type of registration" : 'paid', 
                    "Institution" : 'inst',
                    "City" : 'city',
                    "Country" : 'country'}


def query_yes_no(question, default="yes"):
    """Ask a yes/no question via raw_input() and return their answer.

    "question" is a string that is presented to the user.
    "default" is the presumed answer if the user just hits <Enter>.
        It must be "yes" (the default), "no" or None (meaning
        an answer is required of the user).

    The "answer" return value is True for "yes" or False for "no".
    """
    valid = {"yes": True, "y": True, "ye": True,
             "no": False, "n": False}
    if default is None:
        prompt = " [y/n] "
    elif default == "yes":
        prompt = " [Y/n] "
    elif default == "no":
        prompt = " [y/N] "
    else:
        raise ValueError("invalid default answer: '%s'" % default)

    while True:
        sys.stdout.write(question + prompt)
        sys.stdout.flush()
        choice = raw_input().lower()
        if default is not None and choice == '':
            return valid[default]
        elif choice in valid:
            return valid[choice]
        else:
            sys.stdout.write("Please respond with 'yes' or 'no' "
                             "(or 'y' or 'n').\n")

def find_best_user(row, registered_email):
    print('%s is not found in the main registration! Trying to auto guess...' %email)
    from difflib import SequenceMatcher
    best_ratio  = 0
    last_name   = row['Last Name'].strip().title()
    middle_name = row['Middle Name'].strip()
    first_name  = row['First Name'].strip().title()
    full_name   = '%s %s'%(last_name, first_name)

    for user in registered_email.keys():
        ratio_1 = SequenceMatcher(None, email, registered_email[user]['email']).ratio()
        ratio_2 = SequenceMatcher(None, full_name, registered_email[user]['full_name']).ratio()
        ratio   =  ratio_1 + ratio_2
        if ratio > best_ratio:
            best_user  = user
            best_ratio = ratio
    print('Best match found for %s is user %s' %(email, best_user))
    return best_user



with open(main_registrations_csv, 'rb') as csvfile:

    reader = csv.DictReader(csvfile, delimiter=',', quotechar='"')

    for count, row in enumerate(reader):
   
        user = {}
        user['email'] = row['Email']

        if display:
            print('Handling registration for %s'%user['email'])

        last_name     = row['Last Name'].strip().title()
        middle_name   = row['Middle Name'].strip()
        first_name    = row['First Name'].strip().title()
        user['name']  = ('<b>%s</b> %s %s'%(last_name, first_name, middle_name)).strip()
        user['full_name']   = '%s %s'%(last_name, first_name)
        user['first_name']  = '%s' %first_name
        user['middle_name'] = '%s' %middle_name
        user['last_name']   = '%s' %last_name
                
        date         = row['Submit Date']
        user['date'] = time.strptime(date, '%m/%d/%Y %H:%M:%S')

        reg_nm = row['Reg Fee (Non-Member)']
        reg_f  = row['Reg Fee (Faculty)']
        reg_p  = row['Reg Fee (Postdoc)']
        reg_s  = row['Reg Fee (Student)']
        reg_type = row['Registration Type']

        user['type'] = 'Non member' if len(reg_nm)>0 else ('Faculty' if len(reg_f)>0 else ('Postdoc' if len(reg_p)>0 else 'Student'))
        payment_info = reg_nm+reg_f+reg_p+reg_s

        if user['type'] in ['Faculty', 'Postdoc', 'Student']: 
            members[user['type'].lower()] +=1
        elif reg_type in ['Faculty', 'Postdoc', 'Student']:
            non_members[reg_type.lower()] += 1
        else:
            print('Problem at line %d with user %s %s %s'%(count, user['email'], user['type'], reg_type))
            print('Press f/p/s if you know Faculty/Postdoc/student, q to quit')
            key = ''
            while key not in ['f', 'p', 's', 'q']:
                key = raw_input('').lower()

            if key == "q":
                sys.exit()
            elif key == 's':
                non_members['student'] += 1
                reg_type = 'Student'
            elif key == 'f':
                non_members['faculty'] += 1
                reg_type = 'Faculty'
            elif key == 'p':
                non_members['postdoc'] += 1
                reg_type = 'Postdoc'


        user['type0'] = '%s %s'%(user['type'], reg_type if len(reg_type)>0 else 'Member')
        user['paid']  = ''

        if 'Main Meeting' in payment_info:
            user['paid'] += 'MM '
            meeting['main'] +=1

        if 'Tutorial' in payment_info:
            user['paid'] += 'T '
            meeting['tutorial']+=1

        if 'Workshops 1 Day Only' in payment_info:
            user['paid'] += 'WS1day '  
            meeting['ws1']+=1

        elif 'Workshops' in payment_info:
            user['paid'] += 'WS2day ' 
            meeting['ws2']+=1

        shirt = ''
        for size in ['S', 'M', 'L', 'XL']:
            nb_shirt = row['Shirt {}'.format(size)]
            if int(nb_shirt) > 0: shirt += nb_shirt + ' * {} '.format(size)
            user['shirt'] = shirt
            tshirts[size]+= int(nb_shirt)

        banquet = row['BanquetTickets']
        meal    = row['Special Meal']

        user['banquet']      = '%s'%('' if banquet == '0' else banquet)
        user['meal']         = meal

        inst = row['Institution'].strip()
        city = row['City'].strip()          
            
        country = row['Country'].strip()
        user['location'] = '%s, %s'%(inst, country)
        user['inst']     = inst
        user['country']  = country
        user['city']     = city

        if not country in countries.keys(): 
            countries[country] = 0
        countries[country] +=1

        balance = -1*float(row['Balance'])
        user['balance'] = '' if balance==0 else '${:10.2f}'.format(balance)

        registered_email[user['email']] = user
        registered_fullname[user['full_name']] = user


print('Processed %d Main Registrations' %len(registered_email))

if os.path.exists(add_to_registrations_csv):

    with open(add_to_registrations_csv, 'rb') as csvfile:
        
        reader = csv.DictReader(csvfile, delimiter=',', quotechar='"')

        for row in reader:

            email = row['Email']
            if display:
                print('Add items to registration of %s'%email)
            if registered_email.has_key(email):
                user = registered_email[email]
            else:
                best_user = find_best_user(row, registered_email)
                accept = query_yes_no('Do you want to merge users?')
                if accept:
                    user = registered_email[best_user]
                else:
                    print('Please find a match for user %s in file %s' %(email, add_to_registrations_csv))
                    sys.exit()

            changed = False

            reg_nm = row['Reg Fee (Non-Member)']
            reg_f  = row['Reg Fee (Faculty)']
            reg_p  = row['Reg Fee (Postdoc)']
            reg_s  = row['Reg Fee (Student)']

            payment_info = reg_nm+reg_f+reg_p+reg_s
            
            #if display:
            #    print('Paid: %s'%payment_info)

            curr = user['paid']

            if 'Main Meeting' in payment_info:
                curr += ' & MM '
                meeting['main'] +=1
                changed = True

            if 'Tutorial' in payment_info:
                curr += ' & T '
                meeting['tutorial']+=1
                changed = True

            if 'Workshops 1 Day Only' in payment_info:
                curr += ' & WS1day'
                meeting['ws1'] +=1
                changed = True

            elif 'Workshops' in payment_info:
                curr += ' & WS2day ' 
                meeting['ws2'] +=1
                changed = True

            user['paid'] = curr

            if not changed:
                print('Please check: Nothing changed for {m} at line {l} !!'.format(m=email, l=count))
                #sys.exit()

if os.path.exists(extras_csv):

    with open(extras_csv, 'rb') as csvfile:

        reader = csv.DictReader(csvfile, delimiter=',', quotechar='"')

        for row in reader:

            email = row['Email']
            if display:
                print('Add extras to registration of %s'%email)
                
            if registered_email.has_key(email):
                user = registered_email[email]
            else:
                best_user = find_best_user(row, registered_email)

                accept = query_yes_no('Do you want to merge users?')
                if accept:
                    user = registered_email[best_user]
                else:
                    print('Please find a match for user %s in file %s' %(email, add_to_registrations_csv))
                    sys.exit()


            changed = False
            banq    = int(row['BanquetTickets'])

            if banq > 0:
                curr = 0 if len(user['banquet']) == 0 else int(user['banquet'])
                user['banquet'] = curr + banq
                changed = True
                extras['banquet'] += banq
                #print('Increasing banquet tickets from %d to %d'%(curr, user['banquet']))
                #print('Increasing extra banquet tickets from %d to %d'%(curr, user['extrabanquet']))

            for size in ['S', 'M', 'L', 'XL']:
                nb_shirt = int(row['Shirt {}'.format(size)])
                shirt = ''
                if int(nb_shirt) > 0: shirt += nb_shirt + ' * {} '.format(size)
                user['shirt'] += shirt
                tshirts[size] += int(nb_shirt)
                if int(nb_shirt) > 0:
                    #print("Adding %d shirts of size %s" %(int(nb_shirt), size))
                    changed = True

            if not changed:
                print('Please check: Nothing changed for {m} at line {l} !!'.format(m=email, l=count))



registrations = []
for name in sorted(registered_fullname.keys()):
    user = registered_fullname[name]
    registrations.append(user)
    

balance_paid = []

for reg in registrations:
    if reg['email'] in balance_paid:
        reg['balance'] = 'PAID'

variables = {}
variables['title'] = "CNS %d registrants" %conf_year
variables['total'] = len(registered_email)

for forms in [meeting, extras, tshirts]:
    for key, value in forms.items():
        variables[key] = value

variables['faculty'] = 'Total Faculty:  %i (%i members, %i non members)'%(members['faculty']+non_members['faculty'], members['faculty'], non_members['faculty'])
variables['postdoc'] = 'Total Postdocs: %i (%i members, %i non members)'%(members['postdoc']+non_members['postdoc'], members['postdoc'], non_members['postdoc'])
variables['student'] = 'Total Students: %i (%i members, %i non members)'%(members['student']+non_members['student'], members['student'], non_members['student'])
variables['all_countries'] =  '%i countries represented'%len(countries)
variables['countries']     = [{'name' : key, 'number' : value} for key, value in countries.items()]
variables['registrations'] = registrations

from jinja2 import Template

template = """
<title>{{ title }}</title>

<h2>Total number of registered: {{ total }}</h2>
<p>Total number for main meeting: {{ main }}</p>
<p>Total number for tutorials: {{ tutorial }}</p>
<p>Total number for 2 days of workshops: {{ ws2 }}</p>
<p>Total number for 1 days of workshops: {{ ws1 }}</p>

<br/>
<p> {{ faculty }} </p>
<p> {{ postdoc }} </p>
<p> {{ student }} </p>
<p> {{ all_countries }} </p>
<br/>F

<table border="1">
<th>
    <td>Country</td>
</th>

{% for reg in countries %}
  <tr>
      <td>{{ reg.name }}</td>
      <td>{{ reg.number }}</td>
  </tr>
{% endfor %}
</table>


<p>Number of banquet tickets: {{ banquet }}</p>
<p>Number of shirts S: {{ S }}</p>
<p>Number of shirts M: {{ M }}</p>
<p>Number of shirts L: {{ L }}</p>
<p>Nsumber of shirts XL: {{ XL }}</p>

<table border="1">
<th>
    <td>Email</td>
    <td>Type</td>
    <td>Paid</td>
    <td>Shirt</td>
    <td>Banquet tickets</td>
    <td>To pay</td>
    <td>Special Meal</td>
</th>

{% for reg in registrations %}
  <tr>
    <td>{{ reg.name }}</td>
    <td>{{ reg.email }}</td>
    <td>{{ reg.type0 }}</td>
    <td>{{ reg.paid }}</td>
    <td>{{ reg.shirt }}</td>
    <td>{{ reg.banquet }}</td>
    <td>{{ reg.balance }}</td>
    <td>{{ reg.meal }}</td>
  </tr>
{% endfor %}
</table>
"""
t = Template(template)

print("Rendering...")
contents = unicode(t.render(variables))

#print contents

import codecs
html = codecs.open('Main_%d.html' %conf_year, 'w',encoding='utf8')
html.write(contents)
html.close()



count      = 0
full_csv   = "\t".join(to_write_full.keys()) + "\n"
badges_csv = "\t".join(to_write_badges.keys()) + "\n"

for reg in registrations:
    
    if count == 0:
        table = document.add_table(rows=3, cols=2)

    r = 0 if count in [0,1] else (1 if count in [2,3] else 2)
    c = 0 if count in [0,2,4] else 1

    name = reg['full_name']
    loc = reg['location']

    mycell = table.rows[r].cells[c]
    
    p1 = mycell.add_paragraph()
    r1 = p1.add_run('CNS %d %s' %(conf_year, conf_place))
    r1.bold = True
    r1.font.size = Pt(10)

    p2 = mycell.add_paragraph()
    r2 = p2.add_run(name)
    r2.bold = True
    r2.font.size = Pt(8)
    
    p3 = mycell.add_paragraph()
    r3 = p3.add_run('%s %s'%(reg['type0'], reg['paid']))
    r3.font.size = Pt(8)

    p4 = mycell.add_paragraph()
    r4 = p4.add_run(loc)
    r4.font.size = Pt(6)

    paragraph = mycell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('ocns.png')
        
    full_csv   += "\t".join([str(reg[value]) for value in to_write_full.values()]) + "\n"
    badges_csv += "\t".join([str(reg[value]) for value in to_write_badges.values()]) + "\n"

    count  += 1
    if count == 6:
        #print('Finished page..')
        document.add_page_break()
        count = 0


document.save('badges.docx')

#print badge_csv

file_csv = codecs.open('badges.csv', 'w',encoding='utf8')
file_csv.write(badges_csv)
file_csv.close()

file_csv = codecs.open('fullinfo.csv', 'w',encoding='utf8')
file_csv.write(full_csv)
file_csv.close()

